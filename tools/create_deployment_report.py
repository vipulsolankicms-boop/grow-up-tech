from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Project_GitHub_Pages_Deployment_Report.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GREEN = "E2F0D9"
GRAY = "F2F2F2"
GOLD = "FFF2CC"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_inches: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def add_status_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Area", "Final State", "Why It Matters"]
    widths = [1.8, 1.9, 3.6]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, text, True)
        shade_cell(cell, BLUE)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_width(cell, widths[idx])
    set_repeat_table_header(table.rows[0])

    rows = [
        ("Local dev", "Running on localhost:3000", "Confirmed the Vite app runs locally before changing deployment."),
        ("Build", "npm run build passes", "Verified the production artifact can be generated reliably."),
        ("Security", "0 moderate+ audit issues", "Reduced risk before publishing publicly."),
        ("Pages workflow", "Deploys dist via GitHub Actions", "Prevents raw TSX source from being served by GitHub Pages."),
        ("Custom domain", "Preserved in build artifact", "Ensures GitHub Pages keeps growuptech.co.in after deployment."),
        ("DNS", "Only GitHub Pages A records", "Removed GoDaddy forwarding records that could route traffic elsewhere."),
        ("Runtime", "Live site working", "Confirmed the blank/old page issue was cache/propagation related."),
    ]
    for area, state, why in rows:
        cells = table.add_row().cells
        values = [area, state, why]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, idx == 0)
            set_cell_width(cells[idx], widths[idx])


def add_checklist_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Step", "Command or Setting", "Expected Result"]
    widths = [2.0, 2.8, 2.5]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, text, True)
        shade_cell(cell, BLUE)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_width(cell, widths[idx])
    set_repeat_table_header(table.rows[0])

    rows = [
        ("Install", "npm install or npm ci", "Dependencies install from lockfile."),
        ("Type check", "npm run lint", "TypeScript passes."),
        ("Build", "npm run build", "dist folder is generated."),
        ("Audit", "npm audit --audit-level=moderate", "0 vulnerabilities."),
        ("Workflow source", "Settings -> Pages -> GitHub Actions", "GitHub runs the Pages workflow."),
        ("Artifact", "upload-pages-artifact path: dist", "Only production files deploy."),
        ("Custom domain", "public/CNAME", "dist/CNAME exists after build."),
        ("DNS", "4 GitHub Pages A records", "DNS check successful."),
        ("Browser QA", "Incognito + hard refresh", "No stale redirects or blank page."),
    ]
    for step, command, expected in rows:
        cells = table.add_row().cells
        values = [step, command, expected]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, idx == 0)
            set_cell_width(cells[idx], widths[idx])


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("GitHub Pages Deployment Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Grow Up Tech / MicroC@re Vite React Project")
    subtitle_run.font.size = Pt(13)
    subtitle_run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Purpose: A reusable guide for preparing future Vite projects for GitHub Pages hosting")

    add_callout(
        doc,
        "Executive Summary",
        "The project was moved from a local Vite app and AI Studio-style scaffold to a GitHub Pages-ready static site. The final deployment builds with GitHub Actions, publishes the Vite dist artifact, preserves the custom domain through public/CNAME, and resolves correctly through GitHub Pages DNS records.",
        GREEN,
    )

    doc.add_heading("Final Readiness Snapshot", level=1)
    add_status_table(doc)

    doc.add_heading("Meaningful Actions Completed", level=1)

    doc.add_heading("1. Local Project Startup", level=2)
    doc.add_paragraph("The project was identified as a Vite React application. Dependencies were installed and the dev server was launched locally.")
    add_code_block(doc, "npm install\nnpm run dev\n# App URL: http://localhost:3000")

    doc.add_heading("2. Baseline Validation", level=2)
    doc.add_paragraph("The project was checked with TypeScript, production build, and npm audit before and after fixes. The final command set is listed near the end of this guide.")

    doc.add_heading("3. Runtime and UX Bug Fixes", level=2)
    add_bullets(
        doc,
        [
            "Changed the contact form so Firestore submission errors remain local to the form instead of crashing the whole app.",
            "Added an inline user-facing submit error message.",
            "Replaced the missing repair-service.jpg path with an existing repair station image.",
            "Wired the gallery category buttons to real filtering state.",
            "Updated the browser title and meta description from AI Studio placeholders to branded site metadata.",
        ],
    )

    doc.add_heading("4. Security and Dependency Cleanup", level=2)
    add_bullets(
        doc,
        [
            "Ran npm audit fix and verified zero moderate-or-higher vulnerabilities.",
            "Removed unused server and AI dependencies: @google/genai, dotenv, express, @types/express, and tsx.",
            "Removed unused Gemini API key injection from Vite config to avoid client-side secret exposure risk.",
            "Simplified Firebase initialization and removed unused Firestore error helper code.",
        ],
    )

    doc.add_heading("5. Asset Cleanup", level=2)
    add_bullets(
        doc,
        [
            "Replaced the oversized 6.7 MB logo with a compact SVG and removed the old public asset so it no longer ships in dist.",
            "Checked dist output for stale references such as repair-service.jpg, growuptechlogo.png, and GEMINI_API_KEY.",
        ],
    )

    doc.add_heading("6. GitHub Pages Workflow Fix", level=2)
    doc.add_paragraph("The original Pages workflow uploaded the raw repository. That would serve index.html with a /src/main.tsx reference, which GitHub Pages cannot transpile. The workflow was changed to install dependencies, build Vite, and upload dist.")
    add_code_block(
        doc,
        "steps:\n  - uses: actions/checkout@v4\n  - uses: actions/setup-node@v4\n    with:\n      node-version: 22\n      cache: npm\n  - run: npm ci\n  - run: npm run build\n  - uses: actions/configure-pages@v5\n  - uses: actions/upload-pages-artifact@v3\n    with:\n      path: dist\n  - uses: actions/deploy-pages@v5",
    )

    doc.add_heading("7. Custom Domain Handling", level=2)
    doc.add_paragraph("The root CNAME file was moved into public/CNAME so Vite copies it into dist during the build.")
    add_code_block(doc, "public/CNAME\ngrowuptech.co.in")

    doc.add_heading("8. DNS and GoDaddy Forwarding", level=2)
    doc.add_paragraph("GoDaddy initially had two extra A records controlled by domain forwarding. Those records were locked and could not be edited directly. Removing GoDaddy forwarding automatically removed them.")
    add_callout(
        doc,
        "Correct Apex A Records",
        "185.199.108.153\n185.199.109.153\n185.199.110.153\n185.199.111.153",
        GOLD,
    )

    doc.add_heading("9. Blank Page Debugging", level=2)
    add_bullets(
        doc,
        [
            "Confirmed the GitHub Actions run was successful and produced a Pages artifact.",
            "Verified the production build locally with Vite build output.",
            "Checked for stale asset names and missing public files.",
            "Used browser hard refresh and repeated attempts to clear stale GoDaddy forwarding/browser cache.",
            "Confirmed the site eventually loaded correctly at the custom domain.",
        ],
    )

    doc.add_heading("Reusable Checklist for Future Projects", level=1)
    add_checklist_table(doc)

    doc.add_heading("Common Failure Modes and Fixes", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Symptom", "Likely Cause", "Fix"]
    widths = [2.0, 2.6, 2.7]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        shade_cell(cell, BLUE)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_width(cell, widths[idx])
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Blank page", "JS bundle failed or stale cache", "Open DevTools Console and Network, hard refresh, check /assets/*.js."),
        ("404 on Pages", "Raw repo deployed instead of dist", "Build with npm run build and upload dist artifact."),
        ("Custom domain disappears", "CNAME not in dist", "Put CNAME in public/CNAME."),
        ("DNS check fails", "Wrong or extra A records", "Use only GitHub Pages A records for apex domain."),
        ("A records locked", "Registrar forwarding or product connection", "Remove forwarding/connection first."),
        ("Works locally, not on github.io/repo", "Wrong Vite base path", "Set base to /repo-name/ or use a custom domain."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, idx == 0)
            set_cell_width(cells[idx], widths[idx])

    doc.add_heading("Final Validation Commands", level=1)
    add_code_block(doc, "npm run lint\nnpm run build\nnpm audit --audit-level=moderate")
    add_callout(
        doc,
        "Ready-to-Host Criteria",
        "The project is ready when lint passes, build passes, audit is clean, the Pages workflow deploys successfully, DNS check is successful, and the live site loads in a fresh browser session.",
        GREEN,
    )

    doc.add_heading("Optional Improvements", level=1)
    add_bullets(
        doc,
        [
            "Enable Enforce HTTPS in GitHub Pages after certificate issuance.",
            "Optimize remaining large public images to reduce artifact size.",
            "Enable noUnusedLocals and noUnusedParameters in tsconfig.json.",
            "Add a favicon and richer Open Graph metadata.",
            "Consider code splitting if the JavaScript bundle remains large.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("GitHub Pages Deployment Report - Grow Up Tech / MicroC@re")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
